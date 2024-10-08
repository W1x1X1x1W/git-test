from docx import Document


doc = Document()

doc.add_heading("hallo world", 2)
p =doc.add_paragraph("its so fun to work here \n").add_run().bold = True
p.add_run(" jsjsja \n").bold = True
p.add_run(" anothor one ").italic = True
doc.add_paragraph("hey bitch", style="list bullet" )

doc.save("test_test.docx")