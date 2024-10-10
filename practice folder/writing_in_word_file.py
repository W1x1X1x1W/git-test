from docx import Document
from docx.shared import Pt



doc = Document()
'''add heading text'''

doc.add_heading("first heading", 0)

'''add paragraphs'''
p = doc.add_paragraph("examples of writing changes")
p.add_run("\n this text is bold").bold = True
p.add_run("\n this text is italic").italic = True

doc.add_paragraph("create a bullet 1 \n", style="List Bullet")
doc.add_paragraph("create a bullet 2 \n", style="List Bullet")

'''add paraghraph with diifernt font and size for every paragraph'''
doc.add_paragraph("this is the first paragraph \n")
style = doc.styles["Normal"]
font = style.font
font.name = "MS Gothic"
font.size = Pt(20)


paragraph1 = doc.add_paragraph("..............11\n")
p1 = paragraph1.add_run("this is the second paragraph!!! \n")
p1.font.size = Pt(48)
p1.font.name = "Arial"

paragraph2 = doc.add_paragraph("............22 \n")
p2 = paragraph2.add_run("this is the third paragraph!!!")
p2.font.size = Pt(72)
p2.font.name = "MS Gothic"

'''creating a table and putting information on it'''
head_informations = ["Name", "age", "work"]
some_informations = [
     ["mohammad",22,"coach"],
    ["bashar",44,"teacher"],
    ["ahmed",20,"player"],
    ["bob",80,"job"]
]
table = doc.add_table(rows=1, cols=3, style=doc.styles['OR'])
for i in range(3):
    table.rows[0].cells[i].text = head_informations[i]

for name,age,work in some_informations:
    cell = table.add_row().cells
    cell[0].text = name
    cell[1].text = str(age)
    cell[2].text = work





doc.save("writing.docx")
