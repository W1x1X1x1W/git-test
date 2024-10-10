from docx.api import Document
from docx.shared import Pt


doc = Document("reading.docx")

for p in doc.paragraphs:
    if p.style.name.startswith("Heading") or p.style.name == "Title":
        print(p.text)


for table in doc.tables:
    print("\n new table")
    for row in table.rows :
        print("|".join(cell.text for cell in row.cells))


all_text = ""
for h in doc.paragraphs:
    all_text += f"{h.text} \n"

print(all_text)

read_just_16pt = ""
for m in doc.paragraphs:
        run = m.runs
        if run.font.size == Pt(16):
            read_just_16pt += f"{m.text} \n"

print(read_just_16pt)


